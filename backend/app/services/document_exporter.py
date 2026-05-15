"""
Document Exporter Service.

Converts a Markdown-formatted YAML description document (produced by
DescriptionGenerator) into a downloadable DOCX or PDF file.

Design decisions:
- DOCX: python-docx with explicit style mapping (Heading 1/2/3, Normal, Code).
- PDF:  fpdf2 with manual rendering — pure Python, no system-level deps.
- Mermaid blocks:
    - flowchart TD/graph TD  → parsed and rendered as a colored step-table (DOCX)
      or structured numbered list (PDF) so the flow is visually clear.
    - erDiagram              → rendered as entity field-tables + relationship
      table in both formats.
    - Other Mermaid types    → fall back to styled monospace box.
- Both exporters accept the same (description_text, job_name, metadata) args
  and return a BytesIO stream ready for FastAPI StreamingResponse.
"""

from __future__ import annotations

import re
import textwrap
from io import BytesIO
from typing import Any, Dict, List, Optional


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

# fpdf2's built-in Helvetica/Courier only support latin-1.  LLM-generated text
# routinely contains em-dashes, smart quotes, bullets, etc.  Replace them so
# the PDF renders without FPDFUnicodeEncodingException.
_UNICODE_REPLACEMENTS: List[tuple[str, str]] = [
    ("\u2014", "--"),   # em dash
    ("\u2013", "-"),    # en dash
    ("\u2018", "'"),    # left single quote
    ("\u2019", "'"),    # right single quote
    ("\u201C", '"'),    # left double quote
    ("\u201D", '"'),    # right double quote
    ("\u2022", "-"),    # bullet
    ("\u2026", "..."),  # ellipsis
    ("\u00A9", "(c)"),  # copyright
    ("\u00AE", "(R)"),  # registered
    ("\u2122", "(TM)"), # trademark
    ("\u00D7", "x"),    # multiplication sign
    ("\u2192", "->"),   # right arrow
    ("\u2190", "<-"),   # left arrow
    ("\u2264", "<="),   # less than or equal
    ("\u2265", ">="),   # greater than or equal
    ("\u2260", "!="),   # not equal
]


def _sanitize_for_pdf(text: str) -> str:
    """Replace common Unicode characters with latin-1 safe equivalents."""
    for src, dst in _UNICODE_REPLACEMENTS:
        text = text.replace(src, dst)
    # Strip any remaining non-latin-1 characters rather than crash
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ---------------------------------------------------------------------------
# Mermaid diagram parsers
# ---------------------------------------------------------------------------

# Node shape patterns in priority order (most specific delimiter first)
_MERMAID_NODE_PATTERNS: List[tuple] = [
    (r'\b([A-Za-z_][\w]*)\(\[([^\]]+)\]\)', "terminal"),    # A([label])
    (r'\b([A-Za-z_][\w]*)\(\(([^)]+)\)\)',  "terminal"),    # A((label))
    (r'\b([A-Za-z_][\w]*)\{\{([^}]+)\}\}', "decision"),    # A{{label}}
    (r'\b([A-Za-z_][\w]*)\{([^}]+)\}',     "decision"),    # A{label}
    (r'\b([A-Za-z_][\w]*)\[\[([^\]]+)\]\]', "process"),    # A[[label]]
    (r'\b([A-Za-z_][\w]*)\[([^\]]+)\]',    "process"),     # A[label]
    (r'\b([A-Za-z_][\w]*)>([^\]]+)\]',     "io"),          # A>label]
    (r'\b([A-Za-z_][\w]*)\(([^)]+)\)',     "terminal"),    # A(label)
]


def _parse_mermaid_flow(text: str):
    """
    Parse a Mermaid flowchart / graph block into (nodes, edges).

    nodes : {id: {"label": str, "shape": "process"|"decision"|"terminal"|"io"}}
    edges : [{"from": str, "to": str, "label": str}]
    """
    nodes: Dict[str, dict] = {}
    edges: List[dict] = []

    # Strip header line + comments + subgraph wrappers
    clean = re.sub(r'^\s*(flowchart|graph)\s+\S+\s*$', '', text,
                   flags=re.MULTILINE | re.IGNORECASE)
    clean = re.sub(r'%%[^\n]*', '', clean)
    clean = re.sub(r'\bsubgraph\b[^\n]*', '', clean, flags=re.IGNORECASE)
    clean = re.sub(r'\bend\b', '', clean, flags=re.IGNORECASE)

    # Extract nodes — most-specific bracket type wins
    for pattern, shape in _MERMAID_NODE_PATTERNS:
        for m in re.finditer(pattern, clean, re.DOTALL):
            nid = m.group(1)
            label = re.sub(r'\s+', ' ', m.group(2).strip().strip('"'))
            if nid not in nodes:
                nodes[nid] = {"label": label, "shape": shape}

    # Labeled edges first:  A -->|"label"| B  or  A -- label --> B
    labeled_pairs: set = set()
    _labeled_re = re.compile(
        r'\b([A-Za-z_][\w]*)\s*'
        r'(?:--[->]*|==+[->]*|\.\.+[->]*)\s*'
        r'\|\s*"?([^|"\n]+?)"?\s*\|\s*'
        r'(?:[->]+\s*)?'
        r'([A-Za-z_][\w]*)\b'
    )
    for m in _labeled_re.finditer(clean):
        f, lbl, t = m.group(1), m.group(2).strip(), m.group(3)
        edges.append({"from": f, "to": t, "label": lbl})
        labeled_pairs.add((f, t))
        for nid in (f, t):
            if nid not in nodes:
                nodes[nid] = {"label": nid, "shape": "process"}

    # Plain arrows: A --> B
    for m in re.finditer(r'\b([A-Za-z_][\w]*)\s*--+>\s*([A-Za-z_][\w]*)\b', clean):
        f, t = m.group(1), m.group(2)
        if (f, t) not in labeled_pairs:
            edges.append({"from": f, "to": t, "label": ""})
            for nid in (f, t):
                if nid not in nodes:
                    nodes[nid] = {"label": nid, "shape": "process"}

    # Defensive: if a label ended up as bracket-only artifact, use the node ID
    for nid, node in nodes.items():
        cleaned = re.sub(r'^[\[\]()\{\}\s]+$', '', node["label"])
        if not cleaned:
            node["label"] = nid

    return nodes, edges


def _parse_mermaid_er(text: str):
    """
    Parse a Mermaid erDiagram block into (entities, relationships).

    entities      : {name: [{"type": str, "field": str}]}
    relationships : [{"from": str, "cardinality": str, "to": str, "label": str}]
    """
    entities: Dict[str, list] = {}
    relationships: List[dict] = []

    # Entity blocks: ENTITY_NAME { type field ... }
    for em in re.finditer(r'\b([A-Z_][A-Z0-9_]*)\s*\{([^}]*)\}', text, re.DOTALL):
        ename, body = em.group(1), em.group(2)
        fields: List[dict] = []
        for line in body.strip().splitlines():
            line = line.strip()
            if not line or line.startswith("%%"):
                continue
            parts = line.split(None, 2)
            if len(parts) >= 2:
                fields.append({"type": parts[0], "field": parts[1].rstrip(",")})
        entities[ename] = fields

    # Relationships: A ||--o{ B : "label"
    _rel_re = re.compile(
        r'\b([A-Z_][A-Z0-9_]*)\s+'
        r'([|o{}<]+--[|o{}<]+)\s+'
        r'([A-Z_][A-Z0-9_]*)\s*:\s*"?([^"\n]+)"?'
    )
    for rm in _rel_re.finditer(text):
        relationships.append({
            "from": rm.group(1),
            "cardinality": rm.group(2),
            "to": rm.group(3),
            "label": rm.group(4).strip(),
        })
        for ename in (rm.group(1), rm.group(3)):
            if ename not in entities:
                entities[ename] = []

    return entities, relationships


def _bfs_order(nodes: dict, edges: List[dict]) -> List[str]:
    """Return node IDs in BFS order starting from nodes with no incoming edges."""
    from collections import deque, defaultdict
    if not nodes:
        return []
    incoming: dict = defaultdict(set)
    outgoing: dict = defaultdict(list)
    for e in edges:
        outgoing[e["from"]].append(e["to"])
        incoming[e["to"]].add(e["from"])
    starts = [nid for nid in nodes if nid not in incoming] or [next(iter(nodes))]
    visited: List[str] = []
    seen: set = set()
    q: deque = deque(starts)
    while q:
        nid = q.popleft()
        if nid in seen:
            continue
        seen.add(nid)
        visited.append(nid)
        for child in outgoing.get(nid, []):
            q.append(child)
    for nid in nodes:
        if nid not in seen:
            visited.append(nid)
    return visited


# ---------------------------------------------------------------------------
# DOCX Mermaid visual renderers
# ---------------------------------------------------------------------------

# Node visual styles: (bg_hex, fg_hex, badge_label)
_DOCX_NODE_STYLES: Dict[str, tuple] = {
    "terminal": ("C6EFCE", "1F5C2E", "START / END"),
    "process":  ("DDEEFF", "1F3864", "PROCESS"),
    "decision": ("FFF2CC", "7D4900", "DECISION"),
    "io":       ("FFE8CC", "7D3800", "INPUT / OUTPUT"),
}


def _set_cell_shade(cell: Any, hex_color: str) -> None:
    """Apply background fill to a DOCX table cell via raw XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _remove_cell_borders(cell: Any) -> None:
    """Remove all borders from a DOCX table cell (for decorative-only rows)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _render_mermaid_flowchart_docx(doc: Any, nodes: dict, edges: List[dict]) -> None:
    """
    Render a parsed Mermaid flowchart as a visual 2-column Word table.

    Layout:
      Col 0 (1.2 in) — coloured shape-type badge (PROCESS / DECISION / etc.)
      Col 1 (4.8 in) — full step label; decision nodes also list YES/NO branches
    Arrow rows (no border, just "↓") separate each step.
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ordered = _bfs_order(nodes, edges)
    out_map: Dict[str, List[tuple]] = {}
    edge_to_label: Dict[tuple, str] = {}
    for e in edges:
        out_map.setdefault(e["from"], []).append((e["to"], e["label"]))
        edge_to_label[(e["from"], e["to"])] = e["label"]

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    try:
        tbl.columns[0].width = Inches(1.2)
        tbl.columns[1].width = Inches(4.8)
    except Exception:
        pass

    def _hex_to_rgb(h: str):
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def add_node_row(nid: str) -> None:
        node = nodes.get(nid, {"label": nid, "shape": "process"})
        shape = node["shape"]
        bg, fg, badge = _DOCX_NODE_STYLES.get(shape, ("F0F0F0", "333333", "STEP"))
        fg_rgb = _hex_to_rgb(fg)

        row = tbl.add_row()
        # Badge cell
        c0 = row.cells[0]
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(badge)
        r0.font.bold = True
        r0.font.size = Pt(8)
        r0.font.color.rgb = RGBColor(*fg_rgb)
        _set_cell_shade(c0, bg)

        # Label cell
        c1 = row.cells[1]
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(node["label"])
        r1.font.size = Pt(10)
        r1.font.bold = shape in ("terminal", "decision")
        r1.font.color.rgb = RGBColor(*fg_rgb)
        _set_cell_shade(c1, bg)

        # For decision nodes, list outgoing branches as sub-lines
        if shape == "decision":
            outs = out_map.get(nid, [])
            if outs:
                p2 = c1.add_paragraph()
                branch_texts = []
                for to_id, lbl in outs:
                    dest_label = nodes[to_id]["label"] if to_id in nodes else to_id
                    branch_texts.append(f"  ➜ {lbl or 'Next'}: {dest_label}" if lbl else f"  ➜ {dest_label}")
                r2 = p2.add_run("  |  ".join(branch_texts))
                r2.font.size = Pt(8)
                r2.font.italic = True
                r2.font.color.rgb = RGBColor(*fg_rgb)
                _set_cell_shade(c1, bg)

    def add_arrow_row(label: str = "") -> None:
        row = tbl.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        _remove_cell_borders(c0)
        _remove_cell_borders(c1)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = p0.add_run("↓")
        ar.font.size = Pt(12)
        ar.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        if label:
            p1 = c1.paragraphs[0]
            lr = p1.add_run(f"({label})")
            lr.font.size = Pt(8)
            lr.font.italic = True
            lr.font.color.rgb = RGBColor(0x88, 0x55, 0x00)

    for i, nid in enumerate(ordered):
        if nid not in nodes:
            continue
        if i > 0:
            prev = ordered[i - 1]
            lbl = edge_to_label.get((prev, nid), "")
            add_arrow_row(lbl)
        add_node_row(nid)

    doc.add_paragraph()  # spacer


def _render_mermaid_er_docx(doc: Any, entities: dict, relationships: list) -> None:
    """
    Render a parsed ER diagram as per-entity field tables + a relationship table.
    """
    from docx.shared import Pt, RGBColor

    # --- Entity tables ---
    for ename, fields in entities.items():
        p = doc.add_paragraph()
        r = p.add_run(f"  Entity: {ename}")
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"

        # Header row
        hdr = tbl.rows[0].cells
        for ci, txt in enumerate(("Type", "Field Name")):
            run = hdr[ci].paragraphs[0].add_run(txt)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_shade(hdr[ci], "1F3864")

        if not fields:
            row = tbl.add_row()
            row.cells[0].paragraphs[0].add_run("—").font.size = Pt(9)
            row.cells[1].paragraphs[0].add_run("(no fields)").font.size = Pt(9)
        else:
            for f in fields:
                row = tbl.add_row()
                r0 = row.cells[0].paragraphs[0].add_run(f.get("type", ""))
                r0.font.size = Pt(9)
                r0.font.italic = True
                r1 = row.cells[1].paragraphs[0].add_run(f.get("field", ""))
                r1.font.size = Pt(9)
                r1.font.bold = True
                _set_cell_shade(row.cells[0], "DDEEFF")
                _set_cell_shade(row.cells[1], "F0F4FF")

        doc.add_paragraph()

    # --- Relationships table ---
    if relationships:
        p2 = doc.add_paragraph()
        r2 = p2.add_run("  Entity Relationships")
        r2.font.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        rel_tbl = doc.add_table(rows=1, cols=3)
        rel_tbl.style = "Table Grid"
        for ci, txt in enumerate(("From Entity", "Relationship / Cardinality", "To Entity")):
            run = rel_tbl.rows[0].cells[ci].paragraphs[0].add_run(txt)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_shade(rel_tbl.rows[0].cells[ci], "1F3864")

        for rel in relationships:
            row = rel_tbl.add_row()
            row.cells[0].paragraphs[0].add_run(rel["from"]).font.size = Pt(9)
            row.cells[1].paragraphs[0].add_run(
                f"{rel['cardinality']}  —  {rel['label']}"
            ).font.size = Pt(9)
            row.cells[2].paragraphs[0].add_run(rel["to"]).font.size = Pt(9)
            _set_cell_shade(row.cells[0], "DDEEFF")
            _set_cell_shade(row.cells[1], "FFF9CC")
            _set_cell_shade(row.cells[2], "DDEEFF")

        doc.add_paragraph()


def _render_mermaid_block_docx(doc: Any, block: dict) -> None:
    """
    Route a mermaid block to the appropriate visual renderer.
    Falls back to a styled monospace code box if parsing yields nothing.
    """
    from docx.shared import Pt, RGBColor

    text = block.get("text", "")

    # Flowchart / graph
    if re.search(r'^\s*(flowchart|graph)\s', text, re.MULTILINE | re.IGNORECASE):
        try:
            nodes, edges = _parse_mermaid_flow(text)
            if nodes:
                _render_mermaid_flowchart_docx(doc, nodes, edges)
                return
        except Exception:
            pass

    # ER diagram
    if re.search(r'^\s*erDiagram', text, re.MULTILINE | re.IGNORECASE):
        try:
            entities, relationships = _parse_mermaid_er(text)
            if entities:
                _render_mermaid_er_docx(doc, entities, relationships)
                return
        except Exception:
            pass

    # Fallback: monospace code box
    label_para = doc.add_paragraph()
    run = label_para.add_run("[Diagram — see raw definition below]")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0xAA)
    code_para = doc.add_paragraph(text)
    code_para.style = "Normal"
    for r in code_para.runs:
        r.font.name = "Courier New"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    _apply_paragraph_shade(code_para, "F0F0F0")
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# DOCX exporter (python-docx)
# ---------------------------------------------------------------------------

def export_to_docx(
    description_text: str,
    job_name: str = "Migration Job",
    metadata: Optional[Dict[str, Any]] = None,
) -> BytesIO:
    """
    Convert a Markdown description document to a formatted DOCX file.

    Args:
        description_text: Full Markdown text from DescriptionGenerator.
        job_name: Name of the migration job — used in the document header.
        metadata: Optional dict with extra fields (e.g. job_id, generated_at,
                  llm_model) written into the document cover block.

    Returns:
        BytesIO stream of the .docx file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page layout ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # ── Cover block ─────────────────────────────────────────────────────────
    title_para = doc.add_heading("YAML Business Logic & Technical Description", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Migration Job: {job_name}")
    run.font.size = Pt(13)
    run.font.bold = True

    if metadata:
        sub2 = doc.add_paragraph()
        sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_parts = []
        if metadata.get("job_id"):
            meta_parts.append(f"Job ID: {metadata['job_id']}")
        if metadata.get("generated_at"):
            meta_parts.append(f"Generated: {metadata['generated_at']}")
        if metadata.get("llm_model"):
            meta_parts.append(f"Model: {metadata['llm_model']}")
        run2 = sub2.add_run("  |  ".join(meta_parts))
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()  # spacer

    # ── Parse and render Markdown ────────────────────────────────────────────
    blocks = _parse_markdown_blocks(description_text)
    for block in blocks:
        _render_docx_block(doc, block)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _render_docx_block(doc: Any, block: dict) -> None:
    """Render a single parsed Markdown block into the DOCX document."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    kind = block["kind"]

    if kind == "heading":
        level = min(block["level"], 4)  # python-docx supports 1-4
        para = doc.add_heading(block["text"], level=level)
        # Add visual separator after H1
        if level == 1:
            doc.add_paragraph()

    elif kind == "paragraph":
        para = doc.add_paragraph(block["text"])
        para.style = "Normal"

    elif kind == "bullet_list":
        for item in block["items"]:
            doc.add_paragraph(item, style="List Bullet")

    elif kind == "numbered_list":
        for item in block["items"]:
            doc.add_paragraph(item, style="List Number")

    elif kind == "mermaid":
        # Route to visual renderer (flowchart table or ER diagram table)
        _render_mermaid_block_docx(doc, block)

    elif kind == "code":
        label_para = doc.add_paragraph()
        run = label_para.add_run("[Code]")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0xAA)

        code_para = doc.add_paragraph(block["text"])
        code_para.style = "Normal"
        # Apply monospace font
        for run in code_para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        # Light background shading via XML
        _apply_paragraph_shade(code_para, "F0F0F0")
        doc.add_paragraph()  # spacer after block

    elif kind == "hr":
        pass  # hr lines omitted from Word — headings provide enough separation

    elif kind == "blank":
        pass  # skip empty lines — spacing is handled elsewhere


def _apply_paragraph_shade(para: Any, hex_color: str) -> None:
    """Apply a background shade to a DOCX paragraph via raw XML."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)
    except Exception:
        pass  # non-critical; skip if XML manipulation fails


# ---------------------------------------------------------------------------
# PDF exporter (fpdf2)
# ---------------------------------------------------------------------------

def export_to_pdf(
    description_text: str,
    job_name: str = "Migration Job",
    metadata: Optional[Dict[str, Any]] = None,
) -> BytesIO:
    """
    Convert a Markdown description document to a PDF file.

    Args:
        description_text: Full Markdown text from DescriptionGenerator.
        job_name: Name of the migration job.
        metadata: Optional dict with extra fields written into the cover.

    Returns:
        BytesIO stream of the .pdf file.
    """
    from fpdf import FPDF

    pdf = _MigrationReportPDF(job_name=_sanitize_for_pdf(job_name), metadata=metadata or {})
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Cover block
    safe_name = _sanitize_for_pdf(job_name)
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, "YAML Business Logic & Technical Description", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 13)
    pdf.multi_cell(0, 8, f"Migration Job: {safe_name}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    if metadata:
        parts = []
        if metadata.get("job_id"):
            parts.append(f"Job ID: {metadata['job_id']}")
        if metadata.get("generated_at"):
            parts.append(f"Generated: {metadata['generated_at']}")
        if metadata.get("llm_model"):
            parts.append(f"Model: {_sanitize_for_pdf(str(metadata['llm_model']))}")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 6, "  |  ".join(parts), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)

    pdf.ln(8)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(6)

    # Render blocks
    blocks = _parse_markdown_blocks(description_text)
    for block in blocks:
        _render_pdf_block(pdf, block)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


class _MigrationReportPDF:
    """Thin wrapper that builds an FPDF instance with header/footer."""

    def __new__(cls, job_name: str = "", metadata: dict = None):  # type: ignore[override]
        from fpdf import FPDF

        class _PDF(FPDF):
            _job_name = job_name

            def header(self):
                if self.page_no() == 1:
                    return
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(120, 120, 120)
                self.cell(0, 8, f"YAML Description -- {self._job_name}", align="L", new_x="LMARGIN", new_y="TOP")
                self.cell(0, 8, f"Page {self.page_no()}", align="R", new_x="LMARGIN", new_y="NEXT")
                self.set_text_color(0, 0, 0)
                self.ln(2)

            def footer(self):
                self.set_y(-15)
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(150, 150, 150)
                self.cell(0, 10, "Generated by Legacy Code Migration System - (c) 2026 Technossus", align="C")
                self.set_text_color(0, 0, 0)

        instance = _PDF(orientation="P", unit="mm", format="A4")
        instance.set_margins(left=15, top=15, right=15)
        return instance


def _render_pdf_block(pdf: Any, block: dict) -> None:
    """Render a single parsed Markdown block into the PDF."""
    kind = block["kind"]

    # Sanitize all text fields for latin-1 compatibility
    if "text" in block:
        block = {**block, "text": _sanitize_for_pdf(block["text"])}
    if "items" in block:
        block = {**block, "items": [_sanitize_for_pdf(item) for item in block["items"]]}

    if kind == "heading":
        level = block["level"]
        if level == 1:
            pdf.ln(5)
            pdf.set_font("Helvetica", "B", 15)
            pdf.set_text_color(20, 50, 120)
            pdf.multi_cell(0, 9, block["text"], new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
            # Underline rule
            pdf.set_draw_color(20, 50, 120)
            pdf.line(15, pdf.get_y(), 195, pdf.get_y())
            pdf.set_draw_color(0, 0, 0)
            pdf.ln(3)
        elif level == 2:
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(40, 80, 160)
            pdf.multi_cell(0, 8, block["text"], new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)
        elif level == 3:
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(60, 100, 180)
            pdf.multi_cell(0, 7, block["text"], new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
        else:
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 10)
            pdf.multi_cell(0, 6, block["text"], new_x="LMARGIN", new_y="NEXT")

    elif kind == "paragraph":
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, block["text"], new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    elif kind == "bullet_list":
        pdf.set_font("Helvetica", "", 10)
        for item in block["items"]:
            pdf.multi_cell(0, 6, f"  -  {item}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    elif kind == "numbered_list":
        pdf.set_font("Helvetica", "", 10)
        for idx, item in enumerate(block["items"], start=1):
            pdf.multi_cell(0, 6, f"  {idx}.  {item}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    elif kind == "code":
        pdf.ln(2)
        pdf.set_font("Helvetica", "BI", 8)
        pdf.set_text_color(68, 68, 170)
        pdf.cell(0, 5, "[ Code ]", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Courier", "", 7)
        pdf.set_fill_color(240, 240, 240)
        for line in block["text"].splitlines():
            wrapped = textwrap.wrap(line or " ", width=120) or [" "]
            for wline in wrapped:
                pdf.multi_cell(0, 4.5, wline, fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    elif kind == "mermaid":
        _render_mermaid_block_pdf(pdf, block)

    elif kind == "hr":
        pass  # hr lines omitted from PDF — headings provide enough separation

    elif kind == "blank":
        pass


def _render_mermaid_block_pdf(pdf: Any, block: dict) -> None:
    """
    Render a Mermaid block in the PDF as a structured numbered/bulleted list.

    - flowchart/graph → numbered step list with shape prefix indicators
    - erDiagram       → entity list + relationship list
    - other           → plain grey monospace box fallback
    """
    raw_text = block.get("text", "")

    # --- Flowchart ---
    if re.search(r'^\s*(flowchart|graph)\s', raw_text, re.MULTILINE | re.IGNORECASE):
        try:
            nodes, edges = _parse_mermaid_flow(raw_text)
            if nodes:
                out_map: Dict[str, List[tuple]] = {}
                for e in edges:
                    out_map.setdefault(e["from"], []).append((e["to"], e["label"]))
                ordered = _bfs_order(nodes, edges)

                pdf.ln(2)
                pdf.set_font("Helvetica", "BI", 9)
                pdf.set_text_color(30, 60, 140)
                pdf.cell(0, 5, "[ Process Flow Diagram ]", new_x="LMARGIN", new_y="NEXT")
                pdf.set_text_color(0, 0, 0)
                pdf.ln(1)

                _SHAPE_PREFIX = {
                    "terminal": "(START/END)",
                    "process":  "[STEP]",
                    "decision": "<DECISION>",
                    "io":       ">INPUT/OUTPUT<",
                }
                _SHAPE_COLORS_PDF = {
                    "terminal": (31, 92, 46),
                    "process":  (31, 56, 100),
                    "decision": (125, 73, 0),
                    "io":       (125, 56, 0),
                }

                for i, nid in enumerate(ordered):
                    if nid not in nodes:
                        continue
                    node = nodes[nid]
                    shape = node["shape"]
                    prefix = _SHAPE_PREFIX.get(shape, "[STEP]")
                    color = _SHAPE_COLORS_PDF.get(shape, (0, 0, 0))

                    pdf.set_fill_color(240, 244, 255) if shape == "process" else \
                        pdf.set_fill_color(240, 250, 240) if shape == "terminal" else \
                        pdf.set_fill_color(255, 250, 220) if shape == "decision" else \
                        pdf.set_fill_color(255, 240, 220)

                    pdf.set_font("Helvetica", "B", 8)
                    pdf.set_text_color(*color)
                    step_text = _sanitize_for_pdf(f"  {i+1}. {prefix}  {node['label']}")
                    pdf.multi_cell(0, 6, step_text, fill=True, new_x="LMARGIN", new_y="NEXT")

                    # Show decision branches
                    if shape == "decision":
                        for to_id, lbl in out_map.get(nid, []):
                            dest = nodes[to_id]["label"] if to_id in nodes else to_id
                            branch = _sanitize_for_pdf(
                                f"       -> {lbl or 'Next'}: {dest}" if lbl else f"       -> {dest}"
                            )
                            pdf.set_font("Helvetica", "I", 7)
                            pdf.set_text_color(120, 80, 0)
                            pdf.multi_cell(0, 5, branch, new_x="LMARGIN", new_y="NEXT")

                    pdf.set_text_color(0, 0, 0)

                    # Arrow between steps
                    if i < len(ordered) - 1 and ordered[i+1] in nodes:
                        pdf.set_font("Helvetica", "", 8)
                        pdf.set_text_color(150, 150, 150)
                        pdf.cell(0, 4, "         |", new_x="LMARGIN", new_y="NEXT")
                        pdf.cell(0, 4, "         v", new_x="LMARGIN", new_y="NEXT")
                        pdf.set_text_color(0, 0, 0)

                pdf.ln(3)
                return
        except Exception:
            pass

    # --- ER Diagram ---
    if re.search(r'^\s*erDiagram', raw_text, re.MULTILINE | re.IGNORECASE):
        try:
            entities, relationships = _parse_mermaid_er(raw_text)
            if entities:
                pdf.ln(2)
                pdf.set_font("Helvetica", "BI", 9)
                pdf.set_text_color(30, 60, 140)
                pdf.cell(0, 5, "[ Entity Relationship Diagram ]", new_x="LMARGIN", new_y="NEXT")
                pdf.set_text_color(0, 0, 0)
                pdf.ln(1)

                for ename, fields in entities.items():
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.set_text_color(31, 56, 100)
                    pdf.set_fill_color(221, 238, 255)
                    pdf.multi_cell(0, 7, f"  {_sanitize_for_pdf(ename)}", fill=True,
                                   new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(0, 0, 0)
                    if fields:
                        pdf.set_font("Courier", "", 8)
                        pdf.set_fill_color(240, 244, 255)
                        for f in fields:
                            line = _sanitize_for_pdf(
                                f"    {f.get('type',''):12s}  {f.get('field','')}"
                            )
                            pdf.multi_cell(0, 5, line, fill=True, new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(2)

                if relationships:
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(31, 56, 100)
                    pdf.cell(0, 6, "  Relationships:", new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(0, 0, 0)
                    pdf.set_font("Helvetica", "", 9)
                    for rel in relationships:
                        line = _sanitize_for_pdf(
                            f"    {rel['from']}  {rel['cardinality']}  {rel['to']}"
                            f"  :  {rel['label']}"
                        )
                        pdf.multi_cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")

                pdf.ln(3)
                return
        except Exception:
            pass

    # --- Fallback: monospace grey box ---
    pdf.ln(2)
    pdf.set_font("Helvetica", "BI", 8)
    pdf.set_text_color(68, 68, 170)
    pdf.cell(0, 5, "[ Diagram ]", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Courier", "", 7)
    pdf.set_fill_color(240, 240, 240)
    for line in raw_text.splitlines():
        wrapped = textwrap.wrap(_sanitize_for_pdf(line or " "), width=120) or [" "]
        for wline in wrapped:
            pdf.multi_cell(0, 4.5, wline, fill=True, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)


# ---------------------------------------------------------------------------
# Shared Markdown block parser
# ---------------------------------------------------------------------------

def _parse_markdown_blocks(text: str) -> List[dict]:
    """
    Parse a Markdown document into a flat list of typed blocks.

    Supported block types:
      - heading   : {"kind": "heading",  "level": int, "text": str}
      - paragraph : {"kind": "paragraph", "text": str}
      - bullet_list : {"kind": "bullet_list", "items": List[str]}
      - numbered_list : {"kind": "numbered_list", "items": List[str]}
      - code      : {"kind": "code",  "text": str, "lang": str}
      - mermaid   : {"kind": "mermaid", "text": str}
      - hr        : {"kind": "hr"}
      - blank     : {"kind": "blank"}

    Inline Markdown (bold, italic, links, inline code) is stripped to plain
    text — only block-level structure is preserved for document styling.
    """
    lines = text.splitlines()
    blocks: List[dict] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Fenced code block ─────────────────────────────────────────────
        fence_match = re.match(r"^```(\w*)", line)
        if fence_match:
            lang = fence_match.group(1).lower()
            code_lines: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            kind = "mermaid" if lang == "mermaid" else "code"
            blocks.append({"kind": kind, "text": "\n".join(code_lines), "lang": lang})
            i += 1
            continue

        # ── ATX heading  (#, ##, ###, ####) ──────────────────────────────
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text_raw = _strip_inline(heading_match.group(2).strip())
            blocks.append({"kind": "heading", "level": level, "text": text_raw})
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────
        if re.match(r"^[-*_]{3,}\s*$", line):
            blocks.append({"kind": "hr"})
            i += 1
            continue

        # ── Unordered list ────────────────────────────────────────────────
        if re.match(r"^\s*[-*+]\s+", line):
            items: List[str] = []
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i]):
                items.append(_strip_inline(re.sub(r"^\s*[-*+]\s+", "", lines[i])))
                i += 1
            blocks.append({"kind": "bullet_list", "items": items})
            continue

        # ── Ordered list ──────────────────────────────────────────────────
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(_strip_inline(re.sub(r"^\s*\d+\.\s+", "", lines[i])))
                i += 1
            blocks.append({"kind": "numbered_list", "items": items})
            continue

        # ── Blank line ────────────────────────────────────────────────────
        if line.strip() == "":
            blocks.append({"kind": "blank"})
            i += 1
            continue

        # ── Paragraph (accumulate until blank or special line) ────────────
        para_lines: List[str] = []
        while i < len(lines):
            l = lines[i]
            if (
                l.strip() == ""
                or re.match(r"^#{1,4}\s+", l)
                or re.match(r"^```", l)
                or re.match(r"^[-*_]{3,}\s*$", l)
                or re.match(r"^\s*[-*+]\s+", l)
                or re.match(r"^\s*\d+\.\s+", l)
            ):
                break
            para_lines.append(l)
            i += 1
        if para_lines:
            blocks.append({"kind": "paragraph", "text": _strip_inline(" ".join(para_lines))})
        continue

    return blocks


# Regex to strip inline Markdown: **bold**, *italic*, `code`, [text](url)
_INLINE_RE = re.compile(
    r"\*\*(.+?)\*\*"        # **bold**
    r"|\*(.+?)\*"           # *italic*
    r"|__(.+?)__"           # __bold__
    r"|_(.+?)_"             # _italic_
    r"|`(.+?)`"             # `code`
    r"|\[([^\]]+)\]\([^)]+\)"  # [text](url)
)


def _strip_inline(text: str) -> str:
    """Strip inline Markdown markers, returning plain text."""
    def _replace(m: re.Match) -> str:
        # Return whichever capture group matched
        for g in m.groups():
            if g is not None:
                return g
        return m.group(0)
    return _INLINE_RE.sub(_replace, text).strip()
